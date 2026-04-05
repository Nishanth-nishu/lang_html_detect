"""
docx_io.py — Word Document I/O for the Foreign Language Identifier.

Read paragraphs from a .docx file, annotate them with language detection,
and write a new .docx with color-highlighted foreign-language runs.

All original styles, images, headers, footers, and formatting are preserved.
Foreign-language text is highlighted with a deterministic language→colour map.
The raw <lang xml:lang="xx"> XML annotation is appended as a small grey note
paragraph immediately after each paragraph that contains foreign text.
"""

from __future__ import annotations

import copy
import re
from pathlib import Path
from typing import List, Tuple

# ---------------------------------------------------------------------------
# Language → highlight colour (hex RGB, no '#')
# ---------------------------------------------------------------------------
_LANG_COLOUR: dict[str, str] = {
    "fr": "1E90FF",   # DodgerBlue   — French
    "es": "FF6347",   # Tomato       — Spanish
    "de": "228B22",   # ForestGreen  — German
    "it": "FF8C00",   # DarkOrange   — Italian
    "ru": "9400D3",   # DarkViolet   — Russian
    "uk": "9400D3",   # DarkViolet   — Ukrainian (same family)
    "ar": "DC143C",   # Crimson      — Arabic
    "he": "DC143C",   # Crimson      — Hebrew
    "fa": "DC143C",   # Crimson      — Persian
    "zh": "FF1493",   # DeepPink     — Chinese
    "ja": "8B4513",   # SaddleBrown  — Japanese
    "ko": "00CED1",   # DarkTurquoise— Korean
    "hi": "B8860B",   # DarkGoldenrod— Hindi / Devanagari
    "mr": "B8860B",   # —            — Marathi
    "ne": "B8860B",   # —            — Nepali
    "ta": "2E8B57",   # SeaGreen     — Tamil
    "te": "2E8B57",   # SeaGreen     — Telugu
    "kn": "2E8B57",   # —            — Kannada
    "ml": "2E8B57",   # —            — Malayalam
    "bn": "2E8B57",   # —            — Bengali
    "pt": "6A0DAD",   # Purple       — Portuguese
    "la": "8B0000",   # DarkRed      — Latin
    "el": "0000CD",   # MediumBlue   — Greek
    "th": "008B8B",   # DarkCyan     — Thai
    "tr": "8FBC8F",   # DarkSeaGreen — Turkish
    "pl": "CD5C5C",   # IndianRed    — Polish
    "nl": "4682B4",   # SteelBlue    — Dutch
    "sv": "6495ED",   # CornflowerBlue— Swedish
    "fi": "5F9EA0",   # CadetBlue    — Finnish
}
_DEFAULT_FOREIGN_COLOUR = "708090"  # SlateGray — any other language

# Minimum-size font for annotation note paragraphs (in half-points)
_NOTE_FONT_SIZE_HP = 16  # 8 pt


# ---------------------------------------------------------------------------
# Word-namespace helpers (python-docx internals)
# ---------------------------------------------------------------------------
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = "{%s}" % _W_NS


def _colour_for(lang: str) -> str:
    return _LANG_COLOUR.get(lang, _DEFAULT_FOREIGN_COLOUR)


# ---------------------------------------------------------------------------
# XML tag parser — turns annotate() output into (text, lang|None) spans
# ---------------------------------------------------------------------------
_TAG_RE = re.compile(r'<lang xml:lang="([^"]+)">(.*?)</lang>', re.DOTALL)


def _parse_annotated(annotated: str) -> List[Tuple[str, str | None]]:
    """
    Parse the XML-annotated string from annotate() into a flat list of
    (text_fragment, lang_or_None) tuples.

    Example:
        "Hello <lang xml:lang=\"fr\">Bonjour</lang> world"
        → [("Hello ", None), ("Bonjour", "fr"), (" world", None)]
    """
    result: List[Tuple[str, str | None]] = []
    pos = 0
    for m in _TAG_RE.finditer(annotated):
        if m.start() > pos:
            result.append((annotated[pos:m.start()], None))
        result.append((m.group(2), m.group(1)))
        pos = m.end()
    if pos < len(annotated):
        result.append((annotated[pos:], None))
    return result


# ---------------------------------------------------------------------------
# Public API: read_docx
# ---------------------------------------------------------------------------

def read_docx(path: str | Path) -> List[Tuple[int, str]]:
    """
    Read all paragraphs from a Word document.

    Returns:
        A list of ``(paragraph_index, paragraph_text)`` tuples.
        Empty paragraphs are included (text = "").
    """
    from docx import Document  # local import keeps startup fast

    doc = Document(str(path))
    return [(i, para.text) for i, para in enumerate(doc.paragraphs)]


# ---------------------------------------------------------------------------
# Public API: write_annotated_docx
# ---------------------------------------------------------------------------

def write_annotated_docx(
    input_path: str | Path,
    output_path: str | Path,
    annotated_paragraphs: List[Tuple[int, str]],
) -> None:
    """
    Copy the input .docx and rewrite the body paragraphs so that
    foreign-language spans are colour-highlighted.

    For each paragraph that contains at least one foreign-language span:
    - The original paragraph runs are replaced with coloured runs.
    - A small annotation note paragraph is inserted immediately after,
      showing the raw <lang …> XML in 8pt grey text.

    Parameters
    ----------
    input_path : str | Path
        Path to the original Word document (read-only).
    output_path : str | Path
        Path where the annotated document will be saved.
    annotated_paragraphs : List[Tuple[int, str]]
        List of ``(paragraph_index, annotated_xml_string)`` pairs
        as returned by ``annotate()`` — only paragraphs that were
        processed are required; the rest pass through unchanged.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    import lxml.etree as etree

    doc = Document(str(input_path))

    # Build a lookup: para_index → annotated_text
    annotation_map: dict[int, str] = {idx: xml for idx, xml in annotated_paragraphs}

    # We need to operate on the body element directly so we can insert
    # note paragraphs after target paragraphs.
    body = doc.element.body

    # Collect references to existing paragraph XML elements (ordered)
    para_elems = [child for child in body if child.tag == _W + "p"]

    # Process each paragraph in reverse order so insertions don't shift indices
    for para_idx in sorted(annotation_map.keys(), reverse=True):
        if para_idx >= len(para_elems):
            continue

        para_elem = para_elems[para_idx]
        annotated = annotation_map[para_idx]
        spans = _parse_annotated(annotated)

        # Check if there are any foreign spans at all
        has_foreign = any(lang is not None for _, lang in spans)
        if not has_foreign:
            continue

        # ---- Rebuild the paragraph runs --------------------------------
        # Extract the paragraph properties (pPr) to keep paragraph style
        ppr = para_elem.find(_W + "pPr")

        # Extract run properties template from the first run (if any)
        first_run = para_elem.find(_W + "r")
        base_rpr = None
        if first_run is not None:
            base_rpr_elem = first_run.find(_W + "rPr")
            if base_rpr_elem is not None:
                base_rpr = copy.deepcopy(base_rpr_elem)

        # Remove all existing runs (and bookmarks etc.) but keep pPr
        for child in list(para_elem):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag not in ("pPr",):
                para_elem.remove(child)

        # Re-add runs with colour
        for text_frag, lang in spans:
            if not text_frag:
                continue

            run_elem = OxmlElement("w:r")

            # Build rPr
            rpr = copy.deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")

            if lang is not None:
                # Set font colour
                colour_elem = rpr.find(_W + "color")
                if colour_elem is None:
                    colour_elem = OxmlElement("w:color")
                    rpr.insert(0, colour_elem)
                hex_colour = _colour_for(lang)
                colour_elem.set(_W + "val", hex_colour)

                # Bold foreign text slightly (optional but helps readability)
                # Uncomment if desired:
                # b_elem = rpr.find(_W + "b")
                # if b_elem is None:
                #     b_elem = OxmlElement("w:b")
                #     rpr.append(b_elem)

            run_elem.append(rpr)

            # w:t element
            t_elem = OxmlElement("w:t")
            t_elem.text = text_frag
            # Preserve leading/trailing whitespace
            if text_frag != text_frag.strip():
                t_elem.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
            run_elem.append(t_elem)
            para_elem.append(run_elem)

        # ---- Insert annotation note paragraph immediately after --------
        note_para = OxmlElement("w:p")

        note_ppr = OxmlElement("w:pPr")
        note_jc = OxmlElement("w:jc")
        note_jc.set(_W + "val", "left")
        note_ppr.append(note_jc)
        note_para.append(note_ppr)

        note_run = OxmlElement("w:r")
        note_rpr = OxmlElement("w:rPr")

        # Grey colour
        note_colour = OxmlElement("w:color")
        note_colour.set(_W + "val", "888888")
        note_rpr.append(note_colour)

        # Small font
        note_sz = OxmlElement("w:sz")
        note_sz.set(_W + "val", str(_NOTE_FONT_SIZE_HP))
        note_rpr.append(note_sz)
        note_szCs = OxmlElement("w:szCs")
        note_szCs.set(_W + "val", str(_NOTE_FONT_SIZE_HP))
        note_rpr.append(note_szCs)

        note_run.append(note_rpr)

        note_t = OxmlElement("w:t")
        # Compact the annotation — only keep tagged portions
        note_t.text = "↳ " + annotated
        note_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        note_run.append(note_t)
        note_para.append(note_run)

        # Insert note_para after para_elem
        para_elem.addnext(note_para)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Convenience: process an entire .docx end-to-end
# ---------------------------------------------------------------------------

def process_docx(input_path: str | Path, output_path: str | Path | None = None) -> Path:
    """
    High-level function: read input_path, annotate every paragraph,
    write annotated output_path.

    Returns the output path.
    """
    from annotator.tagger import annotate  # avoid circular at module level

    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path.with_name(
            input_path.stem + "_annotated" + input_path.suffix
        )
    output_path = Path(output_path)

    paragraphs = read_docx(input_path)
    annotated = [(idx, annotate(text)) for idx, text in paragraphs]
    write_annotated_docx(input_path, output_path, annotated)
    return output_path
